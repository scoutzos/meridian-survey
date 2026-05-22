from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BRAND_DARK = "14110d"
BRAND_GOLD = "c9a878"
TEXT = "1a1a1a"
MUTED = "595959"
TABLE_BORDER = "bfbfbf"
TABLE_FILL = "f6f0e6"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, left=120, bottom=80, right=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), TABLE_BORDER)


def set_paragraph_border_bottom(paragraph, color=BRAND_GOLD, sz="6", space="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def style_run(run, font="Inter", size=10.5, color=TEXT, bold=False, italic=False) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def clean_inline(text: str) -> str:
    text = text.strip()
    text = text.replace("\\#", "#").replace("\\_", "_")
    text = text.replace("\\[", "[").replace("\\]", "]")
    text = text.replace("\\.", ".").replace("\\*", "*")
    text = text.replace("\\", "")
    text = re.sub(r"\s{2,}$", "", text)
    return text


INLINE_RE = re.compile(r"(\*\*\*[^*]+?\*\*\*|\*\*[^*]+?\*\*|\*[^*]+?\*)")


def add_inline_runs(paragraph, text: str, *, size=10.5, color=TEXT, font="Inter", italic_default=False) -> None:
    text = clean_inline(text)
    text = text.replace("***Working Default", "**Working Default")
    if text.endswith("*") and text.startswith("**Working Default"):
        text = text[:-1]
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            style_run(run, font=font, size=size, color=color, italic=italic_default)
        token = match.group(0)
        bold = token.startswith("**")
        italic = token.startswith("*") and not token.startswith("**")
        if token.startswith("***"):
            bold = True
            italic = True
            inner = token[3:-3]
        elif token.startswith("**"):
            inner = token[2:-2]
        else:
            inner = token[1:-1]
        run = paragraph.add_run(clean_inline(inner))
        style_run(run, font=font, size=size, color=color, bold=bold, italic=italic or italic_default)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        style_run(run, font=font, size=size, color=color, italic=italic_default)


def body_paragraph(doc: Document, text: str = "", *, justify=True, italic=False, color=TEXT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(p, text, italic_default=italic, color=color)
    return p


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Operating Agreement")
    style_run(r, "Cormorant Garamond", 28, BRAND_DARK, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("of Meridian Collective Group LLC")
    style_run(r, "Cormorant Garamond", 14, BRAND_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("A Georgia Limited Liability Company")
    style_run(r, "Cormorant Garamond", 11, MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("———————")
    style_run(r, "Inter", 11, BRAND_GOLD)

    for text, after in (
        ("3133 Maple Dr., Ste 240 #2333", 3),
        ("Atlanta, Georgia 30305", 18),
        ("Effective Date  ·  _________________, 2026", 12),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        style_run(r, "Inter", 10, TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("WORKING DRAFT")
    style_run(r, "Inter", 9, BRAND_GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENTIAL  ·  FOR THE PARTNERSHIP")
    style_run(r, "Inter", 7, MUTED)
    doc.add_page_break()


def add_label(doc: Document, text: str, *, page_break=False) -> None:
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(clean_inline(text))
    style_run(r, "Inter", 9, BRAND_GOLD, bold=True)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(16)
    set_paragraph_border_bottom(p)
    r = p.add_run(clean_inline(text))
    style_run(r, "Cormorant Garamond", 22, BRAND_DARK)


def add_h2(doc: Document, text: str) -> None:
    text = clean_inline(re.sub(r"^\*\*|\*\*$", "", text).strip())
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    match = re.match(r"^([0-9]+(?:\.[0-9A-Za-z]+)?)(\s+)(.+)$", text)
    if match:
        r = p.add_run(match.group(1) + match.group(2))
        style_run(r, "Inter", 11, BRAND_GOLD, bold=True)
        r = p.add_run(match.group(3))
        style_run(r, "Inter", 11, BRAND_DARK, bold=True)
    else:
        r = p.add_run(text)
        style_run(r, "Inter", 11, BRAND_DARK, bold=True)


def add_reading_guide_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    set_paragraph_border_bottom(p, sz="4", space="6")
    r = p.add_run(clean_inline(text))
    style_run(r, "Cormorant Garamond", 18, BRAND_DARK, italic=True)


def add_signature_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(clean_inline(text))
    style_run(r, "Inter", 10.5, TEXT)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_numbering(p, 1, 0)
    add_inline_runs(p, text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_numbering(p, 2, 0)
    add_inline_runs(p, text)


def set_column_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    set_table_borders(table)
    widths = [1200, 4200, 3960] if col_count == 3 else [9360 // col_count] * col_count
    if rows[0] and rows[0][0].lower() == "member":
        widths = [3100, 4000, 2260]
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_column_width(cell, widths[c_idx])
            if r_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, value, size=9.5, color=TEXT)
            for run in p.runs:
                run.bold = bool(r_idx == 0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip()
        cells = [clean_inline(c.strip()) for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def render_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    try:
        start = next(i for i, line in enumerate(lines) if line.strip() == "**READING GUIDE**")
    except StopIteration:
        start = 0
    i = start
    article_seen = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if line.startswith("# "):
            add_h1(doc, line[2:].strip())
            i += 1
            continue

        if line.startswith("## "):
            add_h2(doc, line[3:].strip())
            i += 1
            continue

        stripped = re.sub(r"^\*\*|\*\*$", "", line)
        if re.fullmatch(r"ARTICLE\s+[IVXLCDM]+", stripped):
            article_seen += 1
            add_label(doc, stripped, page_break=True)
            i += 1
            continue

        if stripped in {"READING GUIDE", "SIGNATURE PAGE", "SCHEDULE A", "SCHEDULE B", "APPENDIX"}:
            add_label(doc, stripped, page_break=stripped != "READING GUIDE")
            i += 1
            continue

        if line in {"*How to read this document*", "*Execution*", "*Initial Capital Contributions*", "*Spousal Acknowledgment Form*", "*Decision Log Summary*"}:
            add_reading_guide_title(doc, line.strip("*"))
            i += 1
            continue

        if line.startswith("* ") and not line.startswith("**"):
            add_bullet(doc, line[2:].strip())
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            add_numbered(doc, numbered.group(1))
            i += 1
            continue

        if set(line) <= {"_", "\\"} and len(line) > 8:
            add_signature_line(doc, "____________________________________________")
            i += 1
            continue

        if line == "———————":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            style_run(r, "Inter", 11, BRAND_GOLD)
            i += 1
            continue

        italic_note = line.startswith("*(") and line.endswith(")*")
        body_paragraph(doc, line, italic=italic_note, color=MUTED if italic_note else TEXT)
        i += 1


def normalize_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["normal"]
    normal.font.name = "Inter"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(TEXT)


def build(template: Path, markdown_path: Path, output: Path) -> None:
    doc = Document(str(template))
    clear_document_body(doc)
    normalize_document(doc)
    add_cover(doc)
    render_markdown_body(doc, markdown_path.read_text(encoding="utf-8"))
    doc.save(str(output))


def main() -> int:
    if len(sys.argv) != 4:
        print("Usage: build-meridian-oa-branded-docx.py TEMPLATE.docx INPUT.md OUTPUT.docx", file=sys.stderr)
        return 2
    build(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
